#!/usr/bin/env python3
"""Generate Word document from commercial-proposal-uk.md"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "commercial-proposal-uk.md"
OUT_PATH = ROOT / "commercial-proposal-uk.docx"


def set_cell_shading(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_formatted_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            paragraph.add_run(part)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def build_document() -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = [parse_table_row(r) for r in table_lines]
            rows = [r for r in rows if not is_separator_row(r)]
            if not rows:
                continue

            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"

            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_formatted_runs(p, cell_text)
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
                        set_cell_shading(cell, "E8EEF4")

            doc.add_paragraph()
            continue

        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                add_formatted_runs(p, re.sub(r"^\d+\.\s", "", lines[i].strip()))
                i += 1
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_runs(p, lines[i].strip()[2:])
                i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Created: {OUT_PATH}")
